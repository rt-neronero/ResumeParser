import logging
import docx
import re
import json
import nltk
from nltk import word_tokenize, pos_tag, ne_chunk
from spacy import load

# Load a pre-trained spaCy model for NER
nlp = load("en_core_web_sm")

logging.basicConfig(level=logging.DEBUG)
class ResumeParser:
    def __init__(self, resume_path):
        self.resume_path = resume_path
        self.resume = {
            'personal_info': {},
            'employment_details': None,
            'education_details': None,
            'skills': None
        }
        self.sections = {
            "personal_info": [
                "Personal Information", "Contact Information", "Contact", "Summary", "Profile", "Overview", "Objective", "Personal Details"
            ],
            "experience": [
                "Experience", "Work Experience", "Professional Experience", "Employment History", "Career History", "Work History",
                "Background Work Experience", "Relevant Experience", "Job Experience", "Project Experience"
            ],
            "education": [
                "Education", "Academic Background", "Academic Qualifications", "Academic Experience", "Education and Training",
                "Educational Background", "Qualifications", "Education History", "Degrees", "Diplomas"
            ],
            "skills": [
                "Skills", "Core Competencies", "Technical Skills", "Professional Skills", "Skill Set", "Competencies", "Abilities", "Key Skills"
            ],
            "certifications": [
                "Certifications", "Licenses", "Certifications and Licenses", "Certificates", "Professional Certifications",
                "Training and Certifications", "Accreditations"
            ],
            "languages": [
                "Languages", "Language Skills", "Spoken Languages", "Languages Known", "Language Proficiency"
            ],
            "additional": [
                "Projects", "Volunteer Experience", "Volunteer Work", "Achievements", "Awards", "Publications",
                "Research", "Professional Development", "Interests", "Hobbies", "Activities", "Affiliations", "Associations" , "Texts in Shapes"
            ],
            "references": [
                "References", "Professional References", "Referees", "Reference Details"
            ]
        }
        self.parse_resume()

    def parse_resume(self):
        # Open the DOCX file and extract the text
        full_text = self._get_full_text()

        logging.debug("\nFull text extracted:\n %s", full_text)

        # Call specific functions to parse sections
        self._parse_personal_info(full_text)
        self._get_experience_section(full_text)
        self._get_education_section(full_text)
        self._get_skills_section(full_text)
    def _get_full_text(self):
        # Open the DOCX file and extract the text
        doc = docx.Document(self.resume_path)
        full_text = '\n'.join([para.text for para in doc.paragraphs])

        # Extract table contents and append to full_text
        for table in doc.tables:
            table_text = self._extract_table_content(table)
            full_text += '\n' + table_text  # Append table content to the full text

        # Extract text from shapes and append to full_text
        full_text += '\n' + self._extract_shapes_text(doc)

        full_text = self._cleanup_text(full_text)
        
        return full_text
    def _extract_table_content(self, table):
        table_text = ""
        
        for cell in self._iter_cells(table):
            # Add text from the current cell
            table_text += cell.text.strip() + "\n" 
        
        return table_text.strip()  # Remove trailing tab/space if any

    def _iter_cells(self, table):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                # Check if the cell is the same as the previous one (either horizontally or vertically merged)
                if (r > 0 and c < len(table.rows[r - 1].cells) and cell._tc is table.rows[r - 1].cells[c]._tc) \
                    or (c > 0 and cell._tc is row.cells[c - 1]._tc):
                    continue
                yield cell
    from lxml import etree

    def _extract_shapes_text(self, doc):
                # Extract the XML tree of the document
        xml_tree = doc.element.getroottree()

        # Define namespaces
        namespaces = {
            'w' : "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        }
        # Extract all text inside <w:t> tags that are inside <v:textbox> tags (which contains actual text content that we want to extract)
        texts = xml_tree.xpath('//w:txbxContent//w:t', namespaces=namespaces)
        # Collect the text content into a list of strings
        extracted_text = [t.text.strip() for t in texts if t.text]
        # retain 1 copy for duplicates
        extracted_text = list(set(extracted_text))
        # Join all the extracted text, ensuring proper formatting and newlines
        formatted_text = '\n'.join(extracted_text)
        # Fix multiple spaces and unwanted newlines, ensuring better readability
        formatted_text = re.sub(r'\n+', '\n', formatted_text)  # Remove excess newlines
        formatted_text = re.sub(r'([^\n])\n([^\n])', r'\1 \2', formatted_text)  # Merge text split across lines
        text = "Texts in Shapes:\n" + formatted_text
        return text

    def _cleanup_text(self, text):
        # Remove unnecessary newlines and multiple blank spaces
        text = re.sub(r'\n+', '\n', text)  # Replace multiple newlines with a single one
        text = text.strip()  # Remove leading and trailing spaces and newlines
        return text
    def _parse_personal_info(self, text):
        # Use regex for better name detection (full name including middle names)
        name_pattern = re.compile(r"([A-Za-z]+(?: (?:[A-Za-z]+\b|[A-Za-z]\.)){0,4} [A-Za-z]+)")  # Handles first, middle, and last names
        name_match = name_pattern.search(text)
        if name_match:
            self.resume['personal_info']['Full Name'] = name_match.group(1).strip()

        # Regex for email extraction (to capture multiple emails)
        email_pattern = re.compile(r"([\w.-]+@[\w.-]+)")
        email_matches = email_pattern.findall(text)
        if email_matches:
            self.resume['personal_info']['Email/s'] = list(set(email_matches))  # Remove duplicates, store emails in a list

        # Regex for phone number extraction (including country code and '09' starting numbers)
        phone_pattern = re.compile(r"(\+?\d{1,2}\s?)?(\(?\d{2,4}\)?[\s-]?)?\(?\d{3,4}\)?[\s-]?\d{3,4}[\s-]?\d{3,4}")
        phone_match = phone_pattern.search(text)
        if phone_match:
            self.resume['personal_info']['Phone Number'] = phone_match.group(0).strip()

        # Regex for links (social media, portfolio, etc.)
        link_pattern = re.compile(r"(?<![\w.@])((https?://)?(?:www\.)?[a-zA-Z0-9-]+\.[a-z]{2,6}(?:/[^\s]*)?)")
        links = link_pattern.findall(text)
        filtered_links = []
        unwanted_words_pattern = re.compile(r"\b(at|offered by|in|on)\b", re.IGNORECASE)
        # Check each link
        for link in links:
            link_text = link[0]
            
            # Look for forbidden words specifically *before* the link, not after.
            preceding_text = text[:text.find(link_text)]
            if not re.search(r"\b(at|offered by|on)\b", preceding_text[-30:]):  # check only up to 30 characters before
                filtered_links.append(link_text)

        self.resume['personal_info']['Links'] = filtered_links

    def _get_experience_section(self, text):
        # Find Experience section
        experience_pattern = self._get_section_pattern('experience')
        
        # Find the start of the Experience section
        start_index_match = experience_pattern.search(text)
        if start_index_match:
            start_index = start_index_match.end()  # Move past the "Experience" section header
            
            # Find the next section header to determine the end of the experience section
            end_index = len(text)
            next_section_headers = self._get_next_section_headers('experience')
            for header in next_section_headers:
                header_index = text.lower().find(header.lower(), start_index)
                if header_index != -1 and header_index < end_index:
                    end_index = header_index

            # Extract the experience text
            experience_text = text[start_index:end_index].strip()
            
            if experience_text:
                logging.debug("\nExperience Section Extracted:\n %s", experience_text)
                self._parse_experience_details(experience_text)
            else:
                self.resume['experience'] = "Experience Section content not found."
        else:
            self.resume['experience'] = "Experience Section not found."
    
    def _parse_experience_details(self, experience_text):
        experiences = []
        
        
    def _get_education_section(self, text):
        # Look for the Education section header
        education_pattern = self._get_section_pattern('education')
        
        # Find the start of the Education section
        start_index_match = education_pattern.search(text)
        if start_index_match:
            start_index = start_index_match.end()  # Move past the "Education" section header
            
            # Look for the next section headers to determine the end of the education section
            end_index = len(text)
            next_section_headers = self._get_next_section_headers('education')
            for header in next_section_headers:
                header_index = text.lower().find(header.lower(), start_index)
                if header_index != -1 and header_index < end_index:
                    end_index = header_index
            
            # Extract the education text
            education_text = text[start_index:end_index].strip()
            
            if education_text:
                logging.debug("\nEducation Section Extracted:\n %s", education_text)
                self._parse_education_details(education_text)
            else:
                self.resume['education_details'] = "Education Section content not found."
        else:
            self.resume['education_details'] = "Education Section not found."

    def _parse_education_details(self, text):
        count = 0
        self.resume['education_details'] = []  
        # Define patterns
        institution_pattern = re.compile(r"\b(?:university|college|institute|academy|school|learning center)\b(?!.*\bdiploma\b)", re.IGNORECASE)
        date_pattern = re.compile(r"\b(?:\d{4}|\d{2}[-/]\d{2}[-/]\d{4})\b")
        degree_pattern = re.compile(r"\b(?:b\.?\s?s\.?|m\.?\s?s\.?|ph\.?d|bs-ms|bachelor|master|associate|certification|diploma)\b", re.IGNORECASE)
        # Iterate over lines
        for line in text.splitlines():
            line = line.strip()  # Remove leading/trailing whitespace

            # Check for institution
            if institution_pattern.search(line):
                count += 1
                self.resume['education_details'].append({
                    'institution': line,
                    'degree': None,
                    'date': None,
                })

            # Check for date and update the latest institution entry with date
            elif date_pattern.search(line) and count > 0:
                    self.resume['education_details'][count - 1]['date'] = line

            # Check for degree and update the latest institution entry with degree
            elif degree_pattern.search(line) and count > 0:
                if not self.resume['education_details'][count - 1]['degree']:
                    self.resume['education_details'][count - 1]['degree']= line 
    
    def _get_skills_section(self, text):
        # Look for the Skills section header
        skills_pattern = self._get_section_pattern('skills')
        
        # Find the start of the Skills section
        start_index_match = skills_pattern.search(text)
        
        if start_index_match:
            start_index = start_index_match.end()  # Move past the "Skills" section header
            
            # Look for the next section headers to determine the end of the skills section
            end_index = len(text)
            next_section_headers = self._get_next_section_headers('skills')
            for header in next_section_headers:
                header_index = text.lower().find(header.lower(), start_index)
                if header_index != -1 and header_index < end_index:
                    end_index = header_index
            
            # Extract the skills text
            skills_text = text[start_index:end_index].strip()
            
            if skills_text:
                logging.debug("\nSkills Section Extracted:\n %s", skills_text)
                self._parse_skills_details(skills_text)
            else:
                self.resume['skills'] = "Skills section content not found." #debugger

        else:
            self.resume['skills'] = "Skills Section not found." #debugger

    def _parse_skills_details(self, text):
        self.resume['skills'] = []
        for line in text.splitlines():
            line = line.strip()  # Remove leading/trailing whitespace
            if line:
                self.resume['skills'].append(line)
        pass
    def get_parsed_data(self):
        return json.dumps(self.resume, indent=4, ensure_ascii=False)

    def _get_section_pattern(self, section_name):
        """
        Constructs a regex pattern to match headers for a given section.
        """
        # Fetch the list of headers for the specified section
        headers = self.sections.get(section_name, [])
        
        # Join the headers with the OR `|` operator and escape any regex special characters
        pattern_str = "|".join([re.escape(header) for header in headers])
        
        # Create a compiled regex pattern that matches the section headers at the start of a line or after a newline
        section_pattern = re.compile(rf"(^|\n)({pattern_str})(\n|:|\s+)", re.IGNORECASE)
        
        return section_pattern
    def _get_next_section_headers(self,current_section):
        # Dictionary of keywords for common resume sections
        # Flatten the list of all sections, excluding the current section
        next_section_headers = [
            header for section, headers in self.sections.items() 
            if section != current_section for header in headers
        ]
        
        return next_section_headers

# Usage
resume_parser = ResumeParser(f"resume1.docx")
parsed_data = resume_parser.get_parsed_data()
print(parsed_data)
